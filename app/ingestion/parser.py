import logging
from abc import ABC, abstractmethod
from pathlib import Path

import fitz  # type: ignore[import-untyped]
from docx import Document

from app.ingestion.models import (
    DocumentMetadata,
    DocumentPage,
    ParsedDocument,
)

logger = logging.getLogger(__name__)


class BaseParser(ABC):
    """Base interface for all document parsers."""

    @abstractmethod
    def parse(self, file_path: Path) -> ParsedDocument:
        raise NotImplementedError


def build_metadata(file_path: Path) -> DocumentMetadata:
    return DocumentMetadata(
        filename=file_path.name,
        file_path=file_path.resolve(),
        extension=file_path.suffix.lower(),
        file_size=file_path.stat().st_size,
    )


class PDFParser(BaseParser):
    def parse(self, file_path: Path) -> ParsedDocument:
        pages: list[DocumentPage] = []
        errors: list[tuple[int, str]] = []

        with fitz.open(file_path) as pdf:
            for index, page in enumerate(pdf, start=1):
                try:
                    text = page.get_text("text")
                except Exception as exc:
                    errors.append((index, str(exc)))
                    logger.warning(
                        "PDF page %d extraction failed in %s: %s",
                        index, file_path.name, exc,
                    )
                    continue

                pages.append(DocumentPage(page_number=index, text=text))

        if not pages and errors:
            raise RuntimeError(
                f"All {len(errors)} pages failed to parse in {file_path.name}"
            )

        if errors:
            logger.warning(
                "PDF %s: %d/%d pages had extraction errors",
                file_path.name, len(errors), len(pages) + len(errors),
            )

        return ParsedDocument(
            pages=pages,
            metadata=build_metadata(file_path),
        )


class DOCXParser(BaseParser):
    def parse(self, file_path: Path) -> ParsedDocument:
        document = Document(str(file_path))

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        return ParsedDocument(
            pages=[
                DocumentPage(
                    page_number=1,
                    text=text,
                )
            ],
            metadata=build_metadata(file_path),
        )


class TextFileParser(BaseParser):
    """Parser for plain-text formats (``.txt``, ``.md``, ``.rst``, etc.).

    Reads the file as UTF-8 with fallback encoding attempts for
    non-UTF-8 files.
    """

    _FALLBACK_ENCODINGS: tuple[str, ...] = ("utf-8", "latin-1", "cp1252")

    def parse(self, file_path: Path) -> ParsedDocument:
        text = self._read_with_fallback(file_path)

        return ParsedDocument(
            pages=[
                DocumentPage(
                    page_number=1,
                    text=text,
                )
            ],
            metadata=build_metadata(file_path),
        )

    def _read_with_fallback(self, file_path: Path) -> str:
        """Try UTF-8 first, then fall back through common encodings."""
        for encoding in self._FALLBACK_ENCODINGS:
            try:
                return file_path.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue

        # Last resort: read as bytes and decode with replacement
        raw = file_path.read_bytes()
        logger.warning(
            "Using replacement-character decoding for %s", file_path.name,
        )
        return raw.decode("utf-8", errors="replace")


# ------------------------------------------------------------------
# backward-compatibility aliases (keeps existing imports working)
# ------------------------------------------------------------------


class TXTParser(TextFileParser):
    """Legacy alias — use ``TextFileParser`` directly."""


class MarkdownParser(TextFileParser):
    """Legacy alias — use ``TextFileParser`` directly."""
